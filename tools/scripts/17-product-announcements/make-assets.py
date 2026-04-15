"""Generate the 5 sample Word documents for demo 17 (Product Release Announcements)."""
from docx import Document
import os

OUT = os.path.join(os.path.dirname(__file__), "assets")
os.makedirs(OUT, exist_ok=True)


def make_doc(filename, title, release_date, audience, owner, overview,
             whats_new, impact, limitations, customer_facing):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Release Date: {release_date}")
    doc.add_paragraph(f"Target Audience: {audience}")
    doc.add_paragraph(f"Owner: {owner}")
    doc.add_heading("Overview", 2)
    doc.add_paragraph(overview)
    doc.add_heading("What's New", 2)
    for item in whats_new:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Business Impact", 2)
    doc.add_paragraph(impact)
    doc.add_heading("Known Limitations", 2)
    for item in limitations:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Customer-Facing Summary", 2)
    doc.add_paragraph(customer_facing)
    doc.save(os.path.join(OUT, filename))
    print(f"  {filename}")


# ── 1. Zava Spark AI Assistant v3.0 — CLEAN control doc ─────────────────────
make_doc(
    filename="Zava-Spark-AI-Assistant-v3.docx",
    title="Zava Spark AI Assistant v3.0 — Product Release Announcement",
    release_date="March 14, 2026",
    audience="Partners, Public",
    owner="Marcus Theil, Product Marketing",
    overview=(
        "Zava Spark AI Assistant v3.0 is a major release that delivers real-time context awareness, "
        "expanded language support, and significantly faster response times. This release completes "
        "the transition to our next-generation inference engine and marks full general availability "
        "for all Enterprise customers."
    ),
    whats_new=[
        "Real-time context awareness: Spark now maintains session context across up to 50 "
        "conversation turns without performance degradation.",
        "Multi-language support: Full support for 14 languages, including Japanese, Arabic, and Portuguese.",
        "40% faster response times compared to v2.x, achieved through our new distributed inference layer.",
        "Configurable safety filters: Administrators can now tune content policies per workspace.",
        "Native integration with Zava DataConnect and FlowBuilder.",
    ],
    impact=(
        "v3.0 positions Zava Spark as the enterprise AI assistant of record for multilingual "
        "organizations. Early adopters in the financial services pilot reported a 28% reduction "
        "in time spent on routine information retrieval tasks."
    ),
    limitations=[
        "Audio and video input not yet supported — text and file upload only.",
        "Safety filter customization requires Global Admin permissions.",
        "Real-time context awareness is limited to text conversations; file-based sessions "
        "reset context after each upload.",
    ],
    customer_facing=(
        "Zava Spark AI Assistant v3.0 is available today for all Enterprise customers. "
        "This release brings real-time context awareness, support for 14 languages, and "
        "response times 40% faster than the previous version. Upgrade through the Zava "
        "Admin Center — no downtime required. Contact your account team for migration assistance."
    ),
)

# ── 2. Zava DataConnect Platform Update — jargon + passive voice ─────────────
make_doc(
    filename="Zava-DataConnect-Platform-Update.docx",
    title="Zava DataConnect Platform Update — Product Release Announcement",
    release_date="February 3, 2026",
    audience="Partners, Public",
    owner="Priya Nambiar, Platform Engineering",
    overview=(
        "The February 2026 DataConnect platform update delivers improvements to the data ingestion "
        "layer, connector reliability, and observability tooling. The update was designed to address "
        "throughput constraints identified in high-volume enterprise deployments."
    ),
    whats_new=[
        "Ingestion throughput increased by 3x through architectural improvements to the core pipeline.",
        "17 new pre-built connectors added, including Salesforce, Oracle HCM, and SAP S/4HANA.",
        "New observability dashboard with per-connector health metrics and latency breakdowns.",
        "Automatic schema drift detection with configurable alerting thresholds.",
        "Retry logic for transient connector failures — no manual intervention required.",
    ],
    impact=(
        "This update resolves the throughput ceiling that affected large-scale deployments. "
        "Customers running more than 500 concurrent data streams will see the most significant gains."
    ),
    limitations=[
        "Schema drift detection is not yet available for binary or unstructured data sources.",
        "The observability dashboard requires the DataConnect Pro tier or above.",
        "SAP S/4HANA connector is in public preview — production use is not yet recommended.",
    ],
    customer_facing=(
        "This update has been engineered to resolve throughput constraints encountered in "
        "high-concurrency, sharded multi-tenant ingestion pipeline architectures. Latency "
        "improvements are realized through the decoupling of the serialization and deserialization "
        "layers within the core ETL orchestration subsystem. Data fidelity is maintained across "
        "schema mutation events via deterministic conflict-resolution heuristics applied at the "
        "connector boundary. Customers operating in high-volume environments are encouraged to "
        "reconfigure their rate-limiting parameters in conjunction with the updated SDK prior to "
        "re-initializing their pipeline topology."
    ),
)

# ── 3. Zava SecureVault Compliance Pack — fear language + tense + we/company ─
make_doc(
    filename="Zava-SecureVault-Compliance-Pack.docx",
    title="Zava SecureVault Compliance Pack — Product Release Announcement",
    release_date="January 20, 2026",
    audience="Partners, Public",
    owner="Dana Whitfield, Security and Compliance",
    overview=(
        "The Zava SecureVault Compliance Pack adds automated policy enforcement, audit trail "
        "generation, and pre-built compliance templates for SOC 2, ISO 27001, and HIPAA. "
        "The pack integrates directly with the existing SecureVault module and requires no "
        "additional infrastructure."
    ),
    whats_new=[
        "Automated policy enforcement for 23 NIST CSF control categories.",
        "Immutable audit trail with tamper-evident logging for all vault operations.",
        "Pre-built compliance templates for SOC 2 Type II, ISO 27001, and HIPAA.",
        "Real-time alerts for policy violations with customizable escalation paths.",
        "One-click compliance evidence packages for auditor export.",
    ],
    impact=(
        "Organizations that previously spent 6-8 weeks preparing for SOC 2 audits reported "
        "completing evidence collection in under 4 hours using the Compliance Pack in our beta program."
    ),
    limitations=[
        "HIPAA template covers technical safeguards only — administrative and physical safeguard "
        "documentation must be provided manually.",
        "Audit trail retention is 7 years maximum; longer retention requires external archival.",
        "Policy enforcement rules cannot be customized in the Standard tier.",
    ],
    customer_facing=(
        "Organizations that fail to maintain continuous compliance documentation risk significant "
        "regulatory penalties, audit failures, and reputational damage that can be difficult to "
        "recover from. Failure to comply with SOC 2 or HIPAA requirements will result in audit "
        "findings that could jeopardize your certification status. We designed the Compliance Pack "
        "because the company understands that audit preparation is a burden most teams cannot absorb. "
        "The company built this so your team does not have to spend weeks assembling evidence by hand. "
        "We believe compliance should be automatic, and the company has made it so. Get started today "
        "before your next audit window opens, or risk being unprepared when it matters most."
    ),
)

# ── 4. Zava FlowBuilder Integration Suite — superlatives + grammar ───────────
make_doc(
    filename="Zava-FlowBuilder-Integration-Suite.docx",
    title="Zava FlowBuilder Integration Suite — Product Release Announcement",
    release_date="March 28, 2026",
    audience="Partners, Public",
    owner="Sam Okoye, Ecosystem Partnerships",
    overview=(
        "The FlowBuilder Integration Suite expands FlowBuilder's no-code automation platform with "
        "40 new third-party connectors, a visual flow debugging tool, and shared flow templates "
        "that teams can publish to a central library."
    ),
    whats_new=[
        "40 new connectors including Zendesk, HubSpot, Jira, Slack, and GitHub.",
        "Visual flow debugger with step-by-step execution tracing.",
        "Team template library: publish and discover flows across your organization.",
        "Conditional branching with up to 10 nested logic layers.",
        "Webhook trigger support for real-time event-driven flows.",
    ],
    impact=(
        "Teams using the integration suite in beta reduced manual handoff steps by an average of 62%. "
        "The template library alone saved teams from rebuilding common flows from scratch."
    ),
    limitations=[
        "Visual debugger is not available for flows with more than 200 steps.",
        "Webhook triggers require a publicly reachable endpoint — local testing requires a tunneling tool.",
        "Template library sharing is limited to users within the same tenant.",
    ],
    customer_facing=(
        "FlowBuilder Integration Suite is the best-in-class solution for enterprise workflow "
        "automation — nothing else on the market comes close to the unmatched flexibility and "
        "industry-leading connector ecosystem we have built. This is truly a revolutionary release "
        "that sets an entirely new standard that competitors simply cannot match. Every organizations "
        "needs to be running this. The 40 new connectors make it effortless to integrate with the "
        "tools your team all ready uses every day. Whether you are in IT, operations, or business "
        "development, this is the most powerful automation platform available, period. The template "
        "library ensure your team can hit the ground running on day one."
    ),
)

# ── 5. Zava InsightHub Analytics Dashboard — buries lead + casual tone ───────
make_doc(
    filename="Zava-InsightHub-Analytics-Dashboard.docx",
    title="Zava InsightHub Analytics Dashboard — Product Release Announcement",
    release_date="February 18, 2026",
    audience="Partners, Public",
    owner="Kenji Mori, Data Products",
    overview=(
        "InsightHub's February update delivers a rebuilt dashboard experience, new AI-powered "
        "anomaly detection, and a natural language query interface for non-technical users."
    ),
    whats_new=[
        "Rebuilt dashboard canvas with drag-and-drop widget layout and real-time data binding.",
        "AI anomaly detection: automatic flagging of statistical outliers in time-series data.",
        "Natural language query: ask questions in plain English and get chart or table responses.",
        "Scheduled report delivery via email or Teams with PDF or CSV export.",
        "Role-based dashboard sharing with view/edit permission controls.",
    ],
    impact=(
        "The natural language query interface reduces dependency on data analysts for routine "
        "reporting requests. In the beta, business users resolved 40% of their own data questions "
        "without analyst involvement."
    ),
    limitations=[
        "Natural language query supports structured data sources only — unstructured or "
        "document-based data is not yet supported.",
        "Anomaly detection requires at least 90 days of historical data to establish a reliable baseline.",
        "Scheduled report delivery is limited to 20 recipients per report.",
    ],
    customer_facing=(
        "So a little background before we get into what's new — InsightHub has been around for "
        "a few years now, and honestly it has come a really long way since the early days. We have "
        "heard a lot of feedback from customers over the past year, and we took all of that seriously. "
        "The team has been heads down working on improvements across the board. There have been a lot "
        "of internal debates about prioritization, but we think we landed in a good place. Anyway, "
        "the data infrastructure has evolved quite a bit since we first launched the product, and we "
        "have made a lot of investments in the underlying platform that you might not notice day-to-day "
        "but really matter for reliability and scale. With all that context in mind, we are pretty "
        "excited to share that InsightHub now has a brand new dashboard builder, AI anomaly detection, "
        "and natural language queries. Pretty cool stuff! You guys are gonna love the new query "
        "interface — just type a question and get an answer. We think it is a big deal."
    ),
)

print("Done.")
